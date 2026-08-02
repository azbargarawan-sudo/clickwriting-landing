# -*- coding: utf-8 -*-
"""בונה את קובץ עבודה מספר 2 לפי הנחיות הקורס 16459 (ערים ועיור)."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HEB = "David"
LAT = "Times New Roman"


def set_rtl_paragraph(p):
    pPr = p._p.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    bidi.set(qn('w:val'), '1')


def style_run(run, size=12, bold=False, superscript=False, latin=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = LAT if latin else HEB
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), LAT if latin else HEB)
    rFonts.set(qn('w:hAnsi'), LAT if latin else HEB)
    rFonts.set(qn('w:cs'), HEB)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size * 2)))
    rPr.append(szCs)
    if bold:
        rPr.append(OxmlElement('w:bCs'))
    if not latin:
        rPr.append(OxmlElement('w:rtl'))
    if superscript:
        run.font.superscript = True


def para(doc, text='', size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         spacing=2.0, space_after=0, space_before=0):
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = spacing
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if text:
        style_run(p.add_run(text), size=size, bold=bold)
    return p


def T(p, text, size=12, bold=False, latin=False):
    style_run(p.add_run(text), size=size, bold=bold, latin=latin)


def N(p, num, size=12):
    """מספר הערת שוליים, אחרי סימן הפיסוק."""
    style_run(p.add_run(str(num)), size=size, superscript=True)


def page_break_before(p):
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement('w:pageBreakBefore'))
    return p

def add_page_number_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl_paragraph(p)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    style_run(run, size=12)
    for el, attrs, txt in (('w:fldChar', {'w:fldCharType': 'begin'}, None),
                           ('w:instrText', {'xml:space': 'preserve'}, ' PAGE '),
                           ('w:fldChar', {'w:fldCharType': 'end'}, None)):
        e = OxmlElement(el)
        for k, v in attrs.items():
            e.set(qn(k), v)
        if txt:
            e.text = txt
        run._r.append(e)


doc = Document()

sec = doc.sections[0]
sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
for attr in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(sec, attr, Cm(2.5))
sec._sectPr.append(OxmlElement('w:rtlGutter'))
add_page_number_footer(sec)

normal = doc.styles['Normal']
normal.font.name = HEB
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn('w:cs'), HEB)
normal.element.rPr.rFonts.set(qn('w:ascii'), LAT)
normal.element.rPr.rFonts.set(qn('w:hAnsi'), LAT)

QUESTION = ("תאר, בהתבסס על ספרות המחקר, את מעורבותן של האוכלוסייה היהודית "
            "והאוכלוסייה הערבית בתהליך העיור בארץ-ישראל בשלהי התקופה "
            "העות'מאנית, והצג נתונים כמותיים המבטאים את היקף השפעתן על הערים.")

# ========================= עמוד 1: דף שער =========================
para(doc, spacing=1.0)
para(doc, 'עבודה מספר 2', size=16, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5, space_after=16)
para(doc, 'שאלת העבודה:', bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5, space_after=6)
para(doc, QUESTION, align=WD_ALIGN_PARAGRAPH.CENTER,
     spacing=1.5, space_after=28)

for label, value in (
    ('שם הקורס:', "ערים ועיור בשלהי התקופה העות'מאנית"),
    ('מספר הקורס:', '16459'),
    ('שם המרצה:', "פרופ' תמיר גורן"),
    ('שם פרטי ומשפחה:', 'שהד בשארה'),
    ('מספר תעודת זהות:', '213181282'),
    ('תאריך הגשה:', '31.7.2026'),
):
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5, space_after=4)
    T(p, label + ' ', bold=True)
    T(p, value)

for _ in range(20):
    para(doc, spacing=1.0)

para(doc,
     'הריני מצהיר/ה ומאשר/ת שלא נעשה שימוש בבינה מלאכותית לצורך הכנת עבודה זו.',
     align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)

# ===================== עמודים 2-3: גוף התשובה =====================
doc.add_page_break()

para(doc, "מעורבות האוכלוסייה היהודית והאוכלוסייה הערבית בתהליך העיור",
     size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5,
     space_after=8)

p = para(doc)
T(p, "תהליך העיור בשלהי התקופה העות'מאנית לא הונע בידי גורם דמוגרפי אחד. "
     "האוכלוסייה היהודית הייתה קטנה בהרבה אך כמעט כולה עירונית, ואילו "
     "האוכלוסייה הערבית הייתה רוב מכריע ורובה כפרית, ובכל זאת סיפקה לערים "
     "את עיקר גידולן, את הנהגתן ואת מרבית פעילותן הכלכלית.")

para(doc, 'מעורבות האוכלוסייה היהודית', bold=True,
     align=WD_ALIGN_PARAGRAPH.RIGHT, spacing=1.5, space_before=6, space_after=2)

p = para(doc)
T(p, "הסימן המובהק של הנוכחות היהודית היה ריכוזה בעיר. בן-אריה הראה כי הגידול "
     "היהודי בירושלים הוא שהפך אותה לעיר הגדולה בארץ: מכ-2,000 יהודים בראשית "
     "המאה לכ-45,000 ערב מלחמת העולם הראשונה, וכבר בשנות השבעים של המאה "
     "התשע-עשרה היוו בה היהודים רוב יחסי.")
N(p, 1)
T(p, " המעורבות התבטאה בראש ובראשונה בבנייה. מהקמת משכנות שאננים ב-1860 ועד "
     "המלחמה נוסדו מחוץ לחומות ירושלים עשרות שכונות, ורובן המכריע יהודיות, "
     "ביוזמת חברות בנייה, כוללים ופילנתרופים. ייחודה של בנייה זו היה במימונה: "
     "היא לא צמחה מן הכלכלה המקומית אלא מהון שהוזרם מחוץ לארץ.")
N(p, 2)

p = para(doc)
T(p, "ביפו לבשה המעורבות אופי אחר. חברת אחוזת בית, שנוסדה ב-1906, קבעה מראש "
     "עקרונות בנייה מחייבים: מגרש שלא יקטן מאלף אמות מרובעות, רחובות רחבים "
     "וסלולים עם מדרכות ותאורה, שטחים לגנים ציבוריים וצנרת מים וביוב. ב-1908 "
     "רכשה 150,000 אמות מרובעות מכרם ג'יבלי, כ-85 דונם, תמורת 135,000 פרנק, "
     "ועליהן קמה ב-1909 השכונה על שישים מגרשיה. כץ מדגיש שהמתכננים לא ראו בה "
     "עיר עצמאית אלא פרבר-גנים של יפו, שתוסיף לשמש מקום מרכזי לתושביה.")
N(p, 3)
T(p, " בחיפה נותרה הנוכחות היהודית קטנה, כחמישה עשר אחוזים מן התושבים.")
N(p, 4)
T(p, " המעורבות הייתה גם מוסדית: במועצת עיריית ירושלים ישבו נציגים יהודים לצד "
     "מוסלמים ונוצרים, ובבחירות תרנ\"ח נמנו בעיר 700 בעלי זכות בחירה מוסלמים, "
     "300 נוצרים ו-200 יהודים.")
N(p, 5)

para(doc, 'מעורבות האוכלוסייה הערבית', bold=True,
     align=WD_ALIGN_PARAGRAPH.RIGHT, spacing=1.5, space_before=6, space_after=2)

p = para(doc)
T(p, "גרוסמן חלק על תדמיתה של אוכלוסייה ערבית קפואה ומדולדלת והראה שהיא גדלה "
     "בקצב ניכר ונעה בתוך הארץ.")
N(p, 6)
T(p, " תנועה זו היא שהזינה את ערי החוף. אליהן נהרו פלאחים מן הכפרים, לצד "
     "מהגרים ממצרים שהגיעו בעקבות שלטון אבראהים פאשא, מן החוראן ומלבנון. "
     "יזבק מראה שצמיחתה של חיפה נשענה כמעט כולה על הגירה ערבית, פנימית "
     "וחיצונית, שהפכה עיירת חוף של כאלף תושבים לעיר מעורבת של כ-22,000 נפש "
     "בעלת רוב מוסלמי ומיעוט נוצרי גדול.")
N(p, 4)

p = para(doc)
T(p, "ההנהגה העירונית הייתה ערבית כמעט לחלוטין. ראשי עיריית ירושלים נמנו כולם "
     "עם המשפחות המוסלמיות הנכבדות, בני חוסייני, ח'אלדי, עלמי ודג'אני, "
     "והתכתובת העירונית נוהלה בערבית עד סוף התקופה.")
N(p, 5)
T(p, " הבעלות על הקרקע סביב יפו הייתה ערבית ברובה, ובכך תחמה את ההתפשטות "
     "היהודית: ב-1907 וב-1908 הוטלו הגבלות חמורות על רכישת קרקע בידי יהודים, "
     "ורכישה ישירה מערבים הייתה כמעט בלתי אפשרית, ולכן נרכש כרם ג'יבלי דווקא "
     "מסוחרי קרקעות יהודים. את הבנייה עצמה ביצעו ברובה פועלים ובעלי מלאכה "
     "ערבים, אף שהחברה הצהירה על כוונה להעסיק יהודים.")
N(p, 3)

para(doc, 'נתונים כמותיים', bold=True,
     align=WD_ALIGN_PARAGRAPH.RIGHT, spacing=1.5, space_before=6, space_after=3)

ROWS = [
    ('העיר', 'ראשית המאה ה-19', 'ערב המלחמה', 'ההרכב ב-1914'),
    ('ירושלים', '8,000-10,000', 'כ-70,000', 'כ-45,000 יהודים, כ-25,000 ערבים'),
    ('יפו', '1,000-1,500', 'כ-45,000', 'כ-15,000 יהודים, היתר ערבים'),
    ('חיפה', 'כ-1,000', 'כ-22,000', 'כ-15% יהודים, היתר ערבים'),
    ('עזה', 'כ-8,000', 'כ-40,000', 'ערבית כמעט כולה'),
    ('שכם', 'כ-7,500', 'כ-25,000', 'ערבית כמעט כולה'),
    ('עכו', 'כ-25,000', 'כ-10,000', 'ערבית כמעט כולה'),
]
table = doc.add_table(rows=len(ROWS), cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table._tbl.tblPr.append(OxmlElement('w:bidiVisual'))
WIDTHS = [Cm(2.2), Cm(3.2), Cm(2.8), Cm(7.8)]
for r, data in enumerate(ROWS):
    for c, val in enumerate(data):
        cell = table.cell(r, c)
        cell.width = WIDTHS[c]
        cp = cell.paragraphs[0]
        set_rtl_paragraph(cp)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.line_spacing = 1.0
        cp.paragraph_format.space_after = Pt(0)
        cp.paragraph_format.space_before = Pt(0)
        style_run(cp.add_run(val), size=10, bold=(r == 0))

p = para(doc, spacing=1.0, space_before=2, space_after=6,
         align=WD_ALIGN_PARAGRAPH.RIGHT)
T(p, "טבלה 1: אומדני אוכלוסייה בערי ארץ-ישראל המרכזיות, על-פי בן-אריה.",
  size=10)
N(p, 1, size=10)

p = para(doc)
T(p, "היהודים מנו ערב המלחמה כ-85,000 נפש, פחות מ-12 אחוזים מאוכלוסיית הארץ, "
     "אך למעלה מ-85 אחוזים מהם ישבו בערים, ומשקלם בשלוש הערים הגדולות היה "
     "גדול לאין ערוך ממשקלם הארצי: כ-63,000 מתוך כ-137,000 תושבי ירושלים, יפו "
     "וחיפה, קרוב למחצית. האוכלוסייה הערבית מנתה כ-700,000 נפש והייתה עירונית "
     "בשיעור של כרבע בלבד, אך בערכים מוחלטים סיפקה את רוב תושבי הערים ואת כלל "
     "אוכלוסייתן של עזה, שכם, חברון ועכו.")

p = para(doc)
T(p, "גרוס מעמיד את קצב הגידול השנתי הממוצע של שש הערים הראשיות בין 1880 "
     "ל-1914 על שלושה אחוזים, ובשתים-עשרה ערים על 2.3 אחוזים. יצוא ההדרים "
     "מיפו, שבתחילתו היה כולו בידי ערבים מקומיים, עלה מרבע מיליון תיבות "
     "באמצע שנות התשעים לכמעט חצי מיליון ב-1903 עד 1905, ולמעלה ממיליון "
     "וחצי ב-1913 וב-1914, וכרבע ממנו היה בסוף התקופה תוצרת פרדסים "
     "יהודיים. בנמל יפו עלו ערכי היצוא "
     "והיבוא פי שלושה בין 1900 ל-1913, בעוד ערכו של הסחר הבין-לאומי כולו עלה "
     "פי שניים בלבד, ועודף היבוא עלה פי ארבעה.")
N(p, 7)
T(p, " השפעתן הייתה אפוא שונה בטיבה. הערבית קבעה את גודלן של הערים, את אופיין "
     "המנהלי והמסחרי ואת בסיס הקרקע שלהן, והיהודית קבעה את קצב גידולן של שלוש "
     "ערים ואת צורתן הבנויה.")

# ================= עמוד 4: רשימת הערות השוליים =================
page_break_before(
    para(doc, 'ספרות המחקר: רשימת הערות השוליים', size=13, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5, space_after=10))

B, R = 1, 0  # bold / regular
NOTES = [
    [("יהושע בן-אריה, ", R),
     ('"שנים-עשר היישובים הגדולים בארץ-ישראל במאה התשע-עשרה", ', R),
     ("קתדרה", B), (", 19 (ניסן תשמ\"א), עמ' 83-143.", R)],
    [("יהושע בן-אריה, ", R), ("עיר בראי תקופה: ירושלים החדשה בראשיתה", B),
     (", ירושלים, יד יצחק בן-צבי, 1988, עמ' 472-499.", R)],
    [("יוסף כץ, ", R),
     ('"חברת \'אחוזת בית\' 1906-1909: הנחת היסודות להקמתה של תל-אביב", ', R),
     ("קתדרה", B), (", 33 (תשמ\"ה), עמ' 161-191.", R)],
    [("מחמוד יזבק, ", R),
     ('"חיפה העות\'מאנית: צמיחה דמוגרפית וריבוי עדתי", ', R),
     ("אופקים בגאוגרפיה", B), (", 73-74 (2009), עמ' 58-76.", R)],
    [("רות קרק, ", R),
     ('"פעילות עיריית ירושלים בסוף התקופה העות\'מאנית", ', R),
     ("קתדרה", B), (", 6 (תשל\"ח), עמ' 74-94.", R)],
    [("דוד גרוסמן, ", R),
     ('"האוכלוסייה הערבית בארץ ישראל בתקופה העות\'מאנית: תדמית ומציאות", ', R),
     ("אופקים בגאוגרפיה", B), (", 68-69 (תשס\"ח), עמ' 6-34.", R)],
    [("נחום גרוס, ", R),
     ('"תמורות כלכליות בארץ-ישראל בסוף התקופה העות\'מאנית", ', R),
     ("קתדרה", B), (", 2 (תשל\"ז), עמ' 111-125.", R)],
]

for i, parts in enumerate(NOTES, start=1):
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = 2.0
    pf.space_after = Pt(2)
    pf.left_indent = Cm(0.8)
    pf.first_line_indent = Cm(-0.8)
    T(p, "{}. ".format(i))
    for txt, bold in parts:
        T(p, txt, bold=bool(bold))

# --- Word דורש את צאצאי pPr/rPr/sectPr בסדר הסכמה ---
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
ORDER = {
    'pPr': ['pStyle', 'keepNext', 'keepLines', 'pageBreakBefore', 'framePr',
            'widowControl', 'numPr', 'suppressLineNumbers', 'pBdr', 'shd',
            'tabs', 'suppressAutoHyphens', 'kinsoku', 'wordWrap',
            'overflowPunct', 'topLinePunct', 'autoSpaceDE', 'autoSpaceDN',
            'bidi', 'adjustRightInd', 'snapToGrid', 'spacing', 'ind',
            'contextualSpacing', 'mirrorIndents', 'suppressOverlap', 'jc',
            'textDirection', 'textAlignment', 'textboxTightWrap', 'outlineLvl',
            'divId', 'cnfStyle', 'rPr', 'sectPr', 'pPrChange'],
    'rPr': ['rStyle', 'rFonts', 'b', 'bCs', 'i', 'iCs', 'caps', 'smallCaps',
            'strike', 'dstrike', 'outline', 'shadow', 'emboss', 'imprint',
            'noProof', 'snapToGrid', 'vanish', 'webHidden', 'color', 'spacing',
            'w', 'kern', 'position', 'sz', 'szCs', 'highlight', 'u', 'effect',
            'bdr', 'shd', 'fitText', 'vertAlign', 'rtl', 'cs', 'em', 'lang',
            'eastAsianLayout', 'specVanish', 'oMath'],
    'sectPr': ['footnotePr', 'endnotePr', 'type', 'pgSz', 'pgMar', 'paperSrc',
               'pgBorders', 'lnNumType', 'pgNumType', 'cols', 'formProt',
               'vAlign', 'noEndnote', 'titlePg', 'textDirection', 'bidi',
               'rtlGutter', 'docGrid', 'printerSettings', 'sectPrChange'],
}
KEEP_FIRST = ('headerReference', 'footerReference')

roots = [doc.element.body, doc.styles.element]
for s in doc.sections:
    roots += [s.header._element, s.footer._element]
for root in roots:
    for tag, order in ORDER.items():
        for el in root.iter(W + tag):
            children = list(el)
            head = [c for c in children if c.tag.split('}')[-1] in KEEP_FIRST]
            rest = [c for c in children if c not in head]
            rest.sort(key=lambda c: order.index(c.tag.split('}')[-1])
                      if c.tag.split('}')[-1] in order else len(order))
            for c in head + rest:
                el.append(c)

OUT = "/home/user/clickwriting-landing/עבודה_2_ערים_ועיור_שהד_בשארה.docx"
doc.save(OUT)
print("saved:", OUT)
