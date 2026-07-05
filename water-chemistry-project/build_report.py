"""
Generates the submission Word document (Water_Chemistry_Report.docx) with the
five required sections and the figures embedded. Run AFTER
carbonate_system_model.py so the PNG figures exist.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

HERE = os.path.dirname(os.path.abspath(__file__))
doc = Document()

# ---- base style ----
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def rtl_line(text, size=12, bold=False, color=(0x1b, 0x3a, 0x53), after=2):
    """Centered Hebrew (right-to-left) header line."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement('w:bidi'))          # mark paragraph RTL
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*color)
    rPr = r._element.get_or_add_rPr()
    rtl = OxmlElement('w:rtl'); rtl.set(qn('w:val'), '1'); rPr.append(rtl)
    return p


# =====================================================================
# 0. University / faculty header  (matches the standard BGU cover format)
# =====================================================================
LOGO = os.path.join(HERE, 'bgu_logo.png')
if os.path.exists(LOGO):
    doc.add_picture(LOGO, width=Inches(2.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_after = Pt(6)
else:
    # placeholder to paste the official emblem into
    ph = doc.add_paragraph()
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ph.add_run('[ paste the official Ben-Gurion University logo here / '
                   'הכניסו כאן את סמל האוניברסיטה ]')
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    rtl_line('אוניברסיטת בן-גוריון בנגב', size=14, bold=True)

# faculty / department / course (the logo already carries the university name)
rtl_line('הפקולטה למדעי ההנדסה', size=12)
rtl_line('המחלקה להנדסה כימית', size=12, after=8)
rtl_line('כימיה של המים בהנדסה כימית סביבתית', size=11,
         color=(0x55, 0x55, 0x55), after=12)


def h(text, size=14, color=(0x1b, 0x3a, 0x53), space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*color)
    return p


def body(text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.italic = italic
    return p


def figure(path, caption):
    if os.path.exists(os.path.join(HERE, path)):
        doc.add_picture(os.path.join(HERE, path), width=Inches(5.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    cap.paragraph_format.space_after = Pt(10)


# =====================================================================
# 1. Title block
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Modeling the Aquatic Carbonate System and Water Acidification')
r.bold = True
r.font.size = Pt(17)
r.font.color.rgb = RGBColor(0x0d, 0x1b, 0x2a)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Water Chemistry course, Ben-Gurion University, Spring 2026')
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(
    'Students:\n'
    'Shada Taha, 213700875, [email]\n'
    'Abdallah Awad, 213041486, [email]\n'
    'Lina Awad, 325348571, [email]\n'
    'Yasmin Nassar, 326486032, [email]\n'
    'Lecturer: Oded Nir').font.size = Pt(10.5)

# =====================================================================
# 2. Background  (<= 70 words)
# =====================================================================
h('1.  Background')
body(
    'The carbonate system (CO₂/H₂CO₃*, HCO₃⁻, CO₃²⁻) '
    'governs the pH and buffering capacity of nearly all natural waters. Dissolving '
    'CO₂ forms carbonic acid and lowers pH, the mechanism of ocean and freshwater '
    'acidification, while the same equilibria decide whether calcium carbonate '
    'precipitates (scaling) or dissolves (corrosion). We model how atmospheric CO₂, '
    'alkalinity, temperature and ionic strength jointly control pH, total inorganic '
    'carbon and calcite saturation.'
)  # ~68 words

# =====================================================================
# 3. Methods
# =====================================================================
h('2.  Methods')
body(
    'We built a code-based equilibrium model (provided as an interactive HTML tool and '
    'an equivalent Python/Matplotlib script). '
    'Inputs: CO₂ partial pressure pCO₂ (open system) or total inorganic carbon '
    'Cₜ (closed system), total alkalinity, temperature, ionic strength I, and total '
    'calcium. Alkalinity and hardness are reported both in meq/L and in the common '
    '“mg/L as CaCO₃” convention (multiply eq/L by 50,000 and mol/L by 100,000, '
    'respectively; e.g. 2 meq/L = 100 mg/L as CaCO₃).'
)
body(
    'Equations. The model uses the two carbonic-acid dissociation constants K₁, K₂, '
    'Henry’s constant Kₕ for CO₂, the water ion product Kᵥ, and the calcite '
    'solubility product Kₛₚ. All five are evaluated as functions of temperature with '
    'the Plummer & Busenberg (1982) polynomials. Non-ideality is treated with the Davies '
    'equation, whose Debye-Hückel A-parameter is computed from the '
    'temperature-dependent dielectric constant of water; this converts thermodynamic '
    'constants into conditional (concentration) constants K₁′, K₂′, '
    'Kᵥ′. The proton condition is expressed through alkalinity:'
)
eq = doc.add_paragraph()
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq.add_run('Alk = [HCO₃⁻] + 2[CO₃²⁻] + [OH⁻] − [H⁺]').italic = True
body(
    'Algorithm. For a given alkalinity the model finds the hydrogen-ion concentration '
    '[H⁺] that satisfies this balance by geometric bisection over pH 1-13 '
    '(200 iterations). In the open system [H₂CO₃*] is fixed by Henry’s law '
    '([H₂CO₃*] = Kₕ·pCO₂) and the other species follow from '
    'K₁′, K₂′; in the closed system the fixed Cₜ is distributed over '
    'the three species via the ionization fractions α₀, α₁, α₂.'
)
body(
    'Outputs. Equilibrium pH (activity scale), the concentrations of H₂CO₃*, '
    'HCO₃⁻ and CO₃²⁻, total inorganic carbon Cₜ, and the calcite '
    'saturation state, reported as the saturation index '
    'SI = log₁₀(a_Ca·a_CO₃ / Kₛₚ) and the saturation ratio '
    'Ω. SI > 0 indicates over-saturation (precipitation potential); SI < 0 indicates '
    'under-saturation (dissolution / corrosivity). We also output the CCPP '
    '(Calcium Carbonate Precipitation Potential): the mass of CaCO₃, in mg/L as '
    'CaCO₃, that must precipitate (+) or dissolve (−) to reach equilibrium '
    '(Ω = 1), a standard index for water stability and corrosion control (for '
    'example, desalinated RO water has a strongly negative CCPP and must be '
    'remineralized before distribution, whereas hard water has a positive CCPP and '
    'tends to scale). The '
    'model was validated against literature constants at 25 °C (pK₁ = 6.35, '
    'pK₂ = 10.33, pKₕ = 1.47, pKₛₚ = 8.48, A = 0.51).'
)
body(
    'Dynamic extension (kinetics + mass transfer + multiple phases). Beyond the '
    'instantaneous equilibrium, the tool integrates a time-dependent scenario that '
    'couples three phases and two rate processes: (i) gas/liquid CO₂ mass transfer, '
    'J = kₗa·(Kₕ·pCO₂,atm − [CO₂(aq)]), which changes Cₜ but not '
    'alkalinity, and (ii) calcite precipitation kinetics, R = kₚ·(Ω − 1), which '
    'removes Ca and consumes alkalinity (Δalk = −2R, ΔCₜ = −R). '
    'The coupled ordinary differential equations are integrated explicitly in time.'
)

# =====================================================================
# 4. Results and discussion  (one figure per group member, <=100 words each)
# =====================================================================
h('3.  Results and Discussion')
body(
    'All curves below use a carbonate-buffered freshwater as the base case '
    '(Alk = 2.0 meq/L, Ca = 0.8 mM, T = 15 °C, I = 0.01 M) unless stated. '
    'Figures 1-4 are the equilibrium results (one per group member); Figure 5 is an '
    'advanced dynamic result showcasing model complexity.', italic=True)

figure('fig1_speciation.png',
       'Figure 1. Bjerrum plot: carbonate speciation as a function of pH.')
body(
    'The relative abundance of the three carbonate species is fixed by pH. '
    'H₂CO₃* dominates below pK₁ (6.35), HCO₃⁻ dominates between the two '
    'pK values, and CO₃²⁻ takes over above pK₂ (10.33); the curves cross '
    'exactly at the pK values, where the two neighbouring species are equal. Around the '
    'circum-neutral pH of most natural waters (7-8.5) bicarbonate carries essentially '
    'all of Cₜ, which is why HCO₃⁻ is the principal contributor to alkalinity '
    'and to the buffering capacity of the system.')

figure('fig2_acidification.png',
       'Figure 2. Acidification: equilibrium pH (red) and Cₜ (blue) versus '
       'atmospheric pCO₂ at constant alkalinity.')
body(
    'Raising pCO₂ drives more CO₂ into solution (Henry’s law); the added '
    'H₂CO₃* dissociates and releases H⁺, so pH falls almost linearly with '
    'log pCO₂, the essence of ocean and freshwater acidification. Because alkalinity '
    'is conserved, the extra acidity is buffered by converting CO₃²⁻ and '
    'HCO₃⁻, and total inorganic carbon Cₜ rises. Moving from the pre-industrial '
    '(280 ppm) to a high-emission 2100 level (~1000 ppm) lowers the modelled pH by several '
    'tenths of a unit, matching observed trends.')

figure('fig3_calcite_SI.png',
       'Figure 3. Calcite saturation index versus pCO₂, showing the '
       'over/under-saturation threshold.')
body(
    'The saturation index falls as pCO₂ rises, because acidification lowers pH and '
    'shifts speciation away from CO₃²⁻, reducing the ion-activity product '
    'a_Ca·a_CO₃. The curve crosses SI = 0 near ≈ 1000 ppm: to the left the water '
    'is over-saturated (green) and calcite tends to precipitate; to the right it becomes '
    'under-saturated (red) and calcite dissolves, i.e. the water turns corrosive to '
    'carbonate minerals and cement. This single threshold links climate-scale CO₂ to '
    'practical concerns of scaling versus corrosion in pipes and aquifers.')

figure('fig5_ionic_strength.png',
       'Figure 4. Effect of ionic strength on the apparent pK values and calcite SI.')
body(
    'Increasing ionic strength lowers activity coefficients (Davies equation), which '
    'lowers the apparent pK₁′ and pK₂′, dissolved ions are effectively '
    '“shielded”, so the acids appear stronger. The divalent CO₃²⁻ is '
    'affected more than the monovalent species, so pK₂′ shifts most. Calcite SI '
    'first drops sharply then levels off, because the doubly-charged Ca²⁺ and '
    'CO₃²⁻ activities are strongly reduced at higher I. This explains why '
    'saline waters can hold much more dissolved carbonate before precipitating, and why '
    'ignoring non-ideality overestimates scaling potential.')

figure('fig6_kinetics.png',
       'Figure 5 (advanced). Dynamic model coupling CO₂ mass transfer and calcite '
       'precipitation kinetics across three phases.')
body(
    'A CO₂-rich groundwater (equilibrated with high soil pCO₂) is exposed to the '
    'atmosphere. First, CO₂ degasses by gas/liquid mass transfer: Cₜ falls and pH '
    'rises steeply while Ca stays constant. When the water crosses into '
    'super-saturation (SI = 0, marked), calcite precipitation kinetics switch on and '
    'Ca drops as CaCO₃ forms, consuming alkalinity. The system self-regulates near '
    'SI ≈ 0 as precipitation tracks continued degassing. This reproduces travertine '
    'and pipe-scale formation, and simultaneously demonstrates kinetics, mass '
    'transfer and multiple phases.')

# =====================================================================
# 5. Conclusions (<= 70 words)
# =====================================================================
h('4.  Conclusions')
body(
    'A single set of carbonate equilibria, solved through the alkalinity balance, '
    'reproduces water acidification, buffering, and the precipitation-dissolution '
    'behaviour of calcite. The tool quantifies how CO₂, alkalinity, temperature and '
    'ionic strength move a water across the SI = 0 line, directly useful for predicting '
    'acidification impacts, corrosion control, carbonate scaling in desalination and '
    'cooling systems, and chemical stabilization of treated water.'
)  # ~66 words

out = os.path.join(HERE, 'Water_Chemistry_Report.docx')
doc.save(out)
print('Wrote', out)
