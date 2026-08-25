# -*- coding: utf-8 -*-
"""בונה את קובץ העבודה מתוך תבנית המכללה: דף שער, מבוא, סקירת ספרות ורשימת מקורות."""
import copy
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import content

TEMPLATE = '_template.docx'
OUT = 'סמינריון - ותק באגף ולגיטימיות - טיוטה.docx'

HEB_FONT = 'David'
LAT_FONT = 'Times New Roman'


# סדר האלמנטים בתוך w:rPr מחייב לפי סכמת OOXML; חריגה ממנו גורמת ל-Word לדחות את הקובץ.
RPR_ORDER = ['w:rStyle', 'w:rFonts', 'w:b', 'w:bCs', 'w:i', 'w:iCs', 'w:caps',
             'w:smallCaps', 'w:strike', 'w:dstrike', 'w:outline', 'w:shadow',
             'w:emboss', 'w:imprint', 'w:noProof', 'w:snapToGrid', 'w:vanish',
             'w:webHidden', 'w:color', 'w:spacing', 'w:w', 'w:kern', 'w:position',
             'w:sz', 'w:szCs', 'w:highlight', 'w:u', 'w:effect', 'w:bdr', 'w:shd',
             'w:fitText', 'w:vertAlign', 'w:rtl', 'w:cs', 'w:em', 'w:lang',
             'w:eastAsianLayout', 'w:specVanish', 'w:oMath']


def _rpr_set(rpr, tag, attrs=None):
    """מוסיף או מעדכן אלמנט בתוך rPr במיקומו הנכון לפי הסכמה."""
    existing = rpr.find(qn(tag))
    if existing is not None:
        el = existing
    else:
        el = OxmlElement(tag)
        idx = RPR_ORDER.index(tag)
        anchor = None
        for child in rpr.iterchildren():
            ctag = child.tag.split('}')[-1]
            full = 'w:' + ctag
            if full in RPR_ORDER and RPR_ORDER.index(full) > idx:
                anchor = child
                break
        if anchor is not None:
            anchor.addprevious(el)
        else:
            rpr.append(el)
    for k, v in (attrs or {}).items():
        el.set(qn(k), v)
    return el


def style_run(run, size=12, bold=False, rtl=True):
    rpr = run._element.get_or_add_rPr()
    _rpr_set(rpr, 'w:rFonts', {'w:ascii': LAT_FONT, 'w:hAnsi': LAT_FONT,
                               'w:cs': HEB_FONT})
    _rpr_set(rpr, 'w:b', {'w:val': '1' if bold else '0'})
    _rpr_set(rpr, 'w:bCs', {'w:val': '1' if bold else '0'})
    _rpr_set(rpr, 'w:sz', {'w:val': str(size * 2)})
    _rpr_set(rpr, 'w:szCs', {'w:val': str(size * 2)})
    if rtl:
        _rpr_set(rpr, 'w:rtl', {'w:val': '1'})


PPR_ORDER = ['w:pStyle', 'w:keepNext', 'w:keepLines', 'w:pageBreakBefore', 'w:framePr',
             'w:widowControl', 'w:numPr', 'w:suppressLineNumbers', 'w:pBdr', 'w:shd',
             'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
             'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
             'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
             'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
             'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap', 'w:outlineLvl',
             'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr', 'w:pPrChange']


def set_bidi(par):
    """מוסיף w:bidi במיקומו הנכון בתוך pPr (הסדר מחייב לפי הסכמה)."""
    ppr = par._element.get_or_add_pPr()
    if ppr.find(qn('w:bidi')) is not None:
        return
    el = OxmlElement('w:bidi')
    idx = PPR_ORDER.index('w:bidi')
    anchor = None
    for child in ppr.iterchildren():
        full = 'w:' + child.tag.split('}')[-1]
        if full in PPR_ORDER and PPR_ORDER.index(full) > idx:
            anchor = child
            break
    if anchor is not None:
        anchor.addprevious(el)
    else:
        ppr.append(el)


def clear_after(doc, keep_index):
    """מוחק את כל הפסקאות/הטבלאות שאחרי אינדקס נתון בגוף המסמך."""
    body = doc.element.body
    children = [c for c in body.iterchildren()
                if c.tag in (qn('w:p'), qn('w:tbl'))]
    for c in children[keep_index:]:
        body.remove(c)


def add_par(doc, text, style='Normal', size=12, bold=False, align=None,
            first_line=True, space_after=None, hanging=False):
    p = doc.add_paragraph(style=style)
    set_bidi(p)
    if align is not None:
        p.alignment = align
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = 2.0
    if hanging:
        pf.left_indent = Cm(1.25)
        pf.first_line_indent = Cm(-1.25)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if _is_latin(text) else WD_ALIGN_PARAGRAPH.RIGHT
    elif first_line:
        pf.first_line_indent = Cm(0.85)
    else:
        pf.first_line_indent = Cm(0)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    run = p.add_run(text)
    style_run(run, size=size, bold=bold, rtl=not _is_latin(text))
    return p


def _is_latin(t):
    for ch in t:
        if '֐' <= ch <= '׿':
            return False
        if 'a' <= ch.lower() <= 'z':
            return True
    return False


def page_break(doc):
    p = doc.add_paragraph()
    set_bidi(p)
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def main():
    doc = docx.Document(TEMPLATE)

    # --- דף שער: מחיקת כל מה שאחרי השער ובנייתו מחדש ---
    clear_after(doc, 21)

    for i, p in enumerate(doc.paragraphs):
        for r in p.runs:
            style_run(r, size=14 if i == 0 else 12, bold=(i == 0),
                      rtl=not _is_latin(p.text))

    ps = doc.paragraphs
    # 0 = הכותרת
    ps[0].text = ''
    r = ps[0].add_run(content.TITLE)
    style_run(r, size=16, bold=True)
    ps[0].paragraph_format.first_line_indent = Cm(0)
    ps[0].paragraph_format.line_spacing = 1.5

    def setp(idx, text, size=12, bold=False):
        p = ps[idx]
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        p.text = ''
        rr = p.add_run(text)
        style_run(rr, size=size, bold=bold, rtl=not _is_latin(text))
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    setp(4, 'עבודה מוגשת במסגרת סמינריון תנאי כליאה של אסירים פליליים בישראל ובעולם')
    setp(5, 'מאת')
    setp(6, 'שם מלא, מספר תעודת זהות')
    setp(7, 'שם מלא, מספר תעודת זהות')
    setp(8, 'בהנחיית: שם המרצה')
    setp(9, 'החוג לקרימינולוגיה')
    setp(11, 'המכללה האקדמית אשקלון')
    setp(17, 'אשקלון\t\t\t\t\t\tאלול תשפ"ו, אוגוסט 2026')

    # --- גוף העבודה ---
    for kind, text in content.BODY:
        if kind == 'h1':
            page_break(doc)
            p = doc.add_paragraph(style='Heading 1')
            set_bidi(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            style_run(p.add_run(text), size=14, bold=True)
        elif kind == 'h2':
            p = doc.add_paragraph(style='Heading 2')
            set_bidi(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(12)
            style_run(p.add_run(text), size=12, bold=True)
        else:
            add_par(doc, text)

    # --- רשימת מקורות ---
    page_break(doc)
    p = doc.add_paragraph(style='Heading 1')
    set_bidi(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    style_run(p.add_run('רשימת מקורות'), size=14, bold=True)

    for ref in content.REFS_HE + content.REFS_EN:
        add_par(doc, ref, hanging=True)

    doc.save(OUT)
    print('saved:', OUT)


if __name__ == '__main__':
    main()
