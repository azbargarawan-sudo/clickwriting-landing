# -*- coding: utf-8 -*-
"""Build the improved project proposal document (Hebrew, RTL)."""
import copy
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HE_FONT = 'David'
LATIN_FONT = 'Times New Roman'
BLUE = RGBColor(0x1F, 0x4E, 0x79)
DARK = RGBColor(0x22, 0x22, 0x22)

def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    return e

def rtl_para(p, align='right'):
    pPr = p._p.get_or_add_pPr()
    if pPr.find(qn('w:bidi')) is None:
        pPr.append(_el('w:bidi'))
    p.alignment = {'right': WD_ALIGN_PARAGRAPH.RIGHT,
                   'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'justify': WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    return p

def fmt_run(r, size=12, bold=False, color=None, italic=False):
    r.font.name = HE_FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color is not None:
        r.font.color.rgb = color
    rPr = r._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = _el('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), LATIN_FONT)
    rFonts.set(qn('w:hAnsi'), LATIN_FONT)
    rFonts.set(qn('w:cs'), HE_FONT)
    if rPr.find(qn('w:rtl')) is None:
        rPr.append(_el('w:rtl'))
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        rPr.append(_el('w:szCs', **{'w:val': str(int(size * 2))}))
    else:
        szCs.set(qn('w:val'), str(int(size * 2)))
    bCs = rPr.find(qn('w:bCs'))
    if bold and bCs is None:
        rPr.append(_el('w:bCs'))
    return r

def para(doc, text='', size=12, bold=False, align='right', color=None,
         space_after=6, space_before=0, style=None, spacing=2.0):
    p = doc.add_paragraph(style=style)
    rtl_para(p, align)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if spacing:
        p.paragraph_format.line_spacing = spacing
    if text:
        r = p.add_run(text)
        fmt_run(r, size=size, bold=bold, color=color)
    return p

def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    rtl_para(p, 'right')
    r = p.add_run(text)
    sizes = {1: 17, 2: 14, 3: 12.5}
    fmt_run(r, size=sizes.get(level, 12), bold=True, color=BLUE)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(doc, text, size=12, bold_prefix=None, marker='•'):
    p = doc.add_paragraph()
    rtl_para(p, 'right')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 2.0
    pf = p.paragraph_format
    pf.right_indent = Cm(0.45)
    fmt_run(p.add_run(marker + ' '), size=size)
    if bold_prefix:
        fmt_run(p.add_run(bold_prefix + ': '), size=size, bold=True)
    fmt_run(p.add_run(text), size=size)
    return p

def caption(doc, text):
    return para(doc, text, size=10, bold=True, align='center',
                color=RGBColor(0x44, 0x44, 0x44), space_after=12, spacing=1.15)

def add_image(doc, path, width_cm, cap=None):
    p = doc.add_paragraph()
    rtl_para(p, 'center')
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if cap:
        caption(doc, cap)

def make_table(doc, rows, cols, widths_cm=None, header_repeat=False):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = t._tbl.tblPr
    if tblPr.find(qn('w:bidiVisual')) is None:
        bidi = _el('w:bidiVisual')
        tblStyle = tblPr.find(qn('w:tblStyle'))
        if tblStyle is not None:
            tblStyle.addnext(bidi)
        else:
            tblPr.insert(0, bidi)
    if widths_cm:
        t.autofit = False
        tblGrid = t._tbl.find(qn('w:tblGrid'))
        if tblGrid is not None:
            for gc, w in zip(tblGrid.findall(qn('w:gridCol')), widths_cm):
                gc.set(qn('w:w'), str(int(w * 567)))
        for row in t.rows:
            for i, w in enumerate(widths_cm):
                row.cells[i].width = Cm(w)
    if header_repeat:
        trPr = t.rows[0]._tr.get_or_add_trPr()
        trPr.append(_el('w:tblHeader'))
    return t

def cell_text(cell, text, size=10.5, bold=False, align='right', shade=None):
    if shade:
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(_el('w:shd', **{'w:val': 'clear', 'w:color': 'auto', 'w:fill': shade}))
    p = cell.paragraphs[0]
    rtl_para(p, align)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    for i, line in enumerate(str(text).split('\n')):
        if i > 0:
            p = cell.add_paragraph()
            rtl_para(p, align)
        fmt_run(p.add_run(line), size=size, bold=bold)

def fill_table(t, data, widths=None, header_shade='D9E2F3', sizes=(10.5, 10.5)):
    for j, cell in enumerate(t.rows[0].cells):
        cell_text(cell, data[0][j], size=sizes[0], bold=True, align='center', shade=header_shade)
    for i, row in enumerate(data[1:], start=1):
        for j, val in enumerate(row):
            cell_text(t.rows[i].cells[j], val, size=sizes[1])

def page_number_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fldChar1 = _el('w:fldChar', **{'w:fldCharType': 'begin'})
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = _el('w:fldChar', **{'w:fldCharType': 'end'})
    r = p.add_run()
    fmt_run(r, size=10)
    r._r.append(fldChar1)
    r2 = p.add_run(); fmt_run(r2, size=10); r2._r.append(instrText)
    r3 = p.add_run(); fmt_run(r3, size=10); r3._r.append(fldChar2)

